import streamlit as st
import boto3
import psycopg2
import requests
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import PyPDF2
import io
import os
from dotenv import load_dotenv
import logging
import pickle
import time  # Add this import
import re  # Add this import
import plotly.express as px
import pandas as pd  # Add this import
from datetime import datetime
import uuid
import json
from streamlit import session_state
from docx import Document  # Add this import

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration from environment variables
S3_BUCKET_NAME = os.getenv('S3_BUCKET_NAME')
S3_FOLDER_NAME = os.getenv('S3_FOLDER_APP2', 'App2_PromptEng/')
AWS_ACCESS_KEY = os.getenv('AWS_ACCESS_KEY')
AWS_SECRET_KEY = os.getenv('AWS_SECRET_KEY')
AWS_REGION = os.getenv('AWS_REGION')
POSTGRES_HOST = os.getenv('POSTGRES_HOST')
POSTGRES_PORT = os.getenv('POSTGRES_PORT')
POSTGRES_DB = os.getenv('POSTGRES_DB')
POSTGRES_USER = os.getenv('POSTGRES_USER')
POSTGRES_PASSWORD = os.getenv('POSTGRES_PASSWORD')
MISTRAL_API_KEY = os.getenv('MISTRAL_API_KEY')
MISTRAL_API_ENDPOINT = "https://api.mistral.ai/v1/chat/completions"
MISTRAL_EMBED_API_ENDPOINT = "https://api.mistral.ai/v1/embeddings"

# Initialize S3 Client
s3_client = boto3.client(
    's3',
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region_name=AWS_REGION
)

# Initialize PostgreSQL Connection
def get_db_connection():
    try:
        conn = psycopg2.connect(
            host=POSTGRES_HOST,
            port=POSTGRES_PORT,
            dbname=POSTGRES_DB,
            user=POSTGRES_USER,
            password=POSTGRES_PASSWORD
        )
        return conn
    except Exception as e:
        logger.error(f"Database connection error: {str(e)}")
        raise

conn = get_db_connection()
cur = conn.cursor()

# Initialize database tables
def initialize_database():
    """Initialize database tables with proper error handling."""
    try:
        # Create tables with IF NOT EXISTS clause
        cur.execute('''
            CREATE TABLE IF NOT EXISTS embeddings (
                id SERIAL PRIMARY KEY,
                chunk TEXT,
                embedding FLOAT[]
            )
        ''')
        
        cur.execute('''
            CREATE TABLE IF NOT EXISTS model_state (
                id INTEGER PRIMARY KEY,
                vectorizer BYTEA,
                last_updated TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Ensure the tables were created successfully
        cur.execute("""
            SELECT EXISTS (
                SELECT FROM information_schema.tables 
                WHERE table_name = 'model_state'
            )
        """)
        model_state_exists = cur.fetchone()[0]
        
        cur.execute("""
            SELECT EXISTS (
                SELECT FROM information_schema.tables 
                WHERE table_name = 'embeddings'
            )
        """)
        embeddings_exists = cur.fetchone()[0]
        
        if not (model_state_exists and embeddings_exists):
            raise Exception("Failed to create required tables")
            
        conn.commit()
        logger.info("Database tables initialized successfully")
        
    except Exception as e:
        logger.error(f"Database initialization error: {str(e)}")
        conn.rollback()
        raise

# Modified vectorizer initialization with persistence
@st.cache_resource
def initialize_vectorizer(vectorizer_type):
    try:
        if (vectorizer_type == "TF-IDF"):
            # Try to load fitted vectorizer from database
            cur.execute("SELECT vectorizer FROM model_state WHERE id = 1")
            result = cur.fetchone()
            if result and result[0]:
                vectorizer = pickle.loads(result[0])
                st.success("Loaded previously fitted TF-IDF vectorizer")
                return vectorizer, True
            else:
                vectorizer = TfidfVectorizer(
                    max_features=100,
                    stop_words='english',
                    lowercase=True,
                    dtype=np.float32
                )
                return vectorizer, False
        else:
            # For Mistral-Embed, no need to load from database
            return None, False
    except Exception as e:
        st.error(f"Error initializing vectorizer: {str(e)}")
        if vectorizer_type == "TF-IDF":
            vectorizer = TfidfVectorizer(
                max_features=100,
                stop_words='english',
                lowercase=True,
                dtype=np.float32
            )
            return vectorizer, False
        else:
            return None, False

def extract_text_from_pdf(pdf_content):
    """Extract text from PDF binary content."""
    try:
        pdf_file = io.BytesIO(pdf_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return None

def extract_text_from_docx(docx_content):
    """Extract text from DOCX binary content."""
    try:
        doc = Document(io.BytesIO(docx_content))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return '\n'.join(text)  # Fixed the string literal syntax error here
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return None

def load_document_from_s3(bucket_name, document_key):
    """Load document from S3 and handle multiple document types."""
    try:
        response = s3_client.get_object(Bucket=bucket_name, Key=document_key)
        document_content = response['Body'].read()
        
        # Log document details for debugging
        logger.info(f"Loading document: {document_key}, Size: {len(document_content)} bytes")
        
        file_extension = document_key.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return extract_text_from_pdf(document_content)
        elif file_extension in ['docx', 'doc']:
            return extract_text_from_docx(document_content)
        else:
            try:
                return document_content.decode('utf-8')
            except UnicodeDecodeError:
                return document_content.decode('latin-1')
    except Exception as e:
        logger.error(f"Error loading document from S3: {str(e)}")
        st.error(f"Error loading document {document_key}: {str(e)}")
        return None

def chunk_document_fixed_size(document, chunk_size=512, overlap=50):
    """Fixed-size chunking with sliding window."""
    if not document:
        return []
        
    chunks = []
    words = document.split()
    
    chunk_size_words = chunk_size // 4  # Approximate words per chunk
    for i in range(0, len(words), chunk_size_words - overlap):
        chunk = ' '.join(words[i:i + chunk_size_words])
        if len(chunk) > 8000:
            chunk = chunk[:8000]
        chunks.append(chunk)
    return chunks

def chunk_document_sentence(document, max_chunk_size=8000):
    """Sentence-based chunking preserving semantic meaning."""
    if not document:
        return []
    
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', document)
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        sentence_size = len(sentence)
        
        # If single sentence exceeds max size, split it into smaller chunks
        if sentence_size > max_chunk_size:
            if current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_size = 0
            
            # Split long sentence
            words = sentence.split()
            temp_chunk = []
            temp_size = 0
            
            for word in words:
                if temp_size + len(word) + 1 > max_chunk_size:
                    chunks.append(' '.join(temp_chunk))
                    temp_chunk = [word]
                    temp_size = len(word)
                else:
                    temp_chunk.append(word)
                    temp_size += len(word) + 1
            
            if temp_chunk:
                chunks.append(' '.join(temp_chunk))
            
        # If adding sentence exceeds limit, create new chunk
        elif current_size + sentence_size + 1 > max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [sentence]
            current_size = sentence_size
        else:
            current_chunk.append(sentence)
            current_size += sentence_size + 1
    
    # Add the last chunk
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def chunk_document(document, chunking_strategy, chunk_size=512, overlap=50):
    """Chunk document using selected strategy."""
    if chunking_strategy == "Fixed-Size":
        return chunk_document_fixed_size(document, chunk_size, overlap)
    else:
        return chunk_document_sentence(document)

def call_mistral_embed_api(texts):
    """Call Mistral Embed API to get embeddings."""
    try:
        # Ensure texts are non-empty and properly formatted
        texts = [text[:8000] for text in texts if text.strip()]  # Limit text length to 8000 chars
        if not texts:
            raise ValueError("No valid texts provided for embedding.")
        
        # Process in smaller batches with delay between calls
        batch_size = 5  # Reduced batch size
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            
            headers = {
                "Authorization": f"Bearer {MISTRAL_API_KEY}",
                "Content-Type": "application/json"
            }
            data = {
                "input": batch[0] if len(batch) == 1 else batch,
                "model": "mistral-embed"
            }
            
            # Log the request payload for debugging
            logger.info(f"Request payload for batch {i//batch_size + 1}: {data}")
            
            max_retries = 3
            retry_delay = 5  # seconds
            
            for retry in range(max_retries):
                try:
                    response = requests.post(MISTRAL_EMBED_API_ENDPOINT, headers=headers, json=data)
                    
                    if response.status_code == 429:  # Too Many Requests
                        wait_time = int(response.headers.get('Retry-After', retry_delay * (retry + 1)))
                        logger.warning(f"Rate limit hit, waiting {wait_time} seconds...")
                        time.sleep(wait_time)
                        continue
                    
                    response.raise_for_status()
                    embeddings = response.json().get("data", [])
                    all_embeddings.extend([e.get("embedding") for e in embeddings])
                    
                    # Log success for this batch
                    logger.info(f"Successfully processed batch {i//batch_size + 1}")
                    break  # Success, exit retry loop
                    
                except requests.exceptions.HTTPError as http_err:
                    if response.status_code != 429 or retry == max_retries - 1:
                        raise
                
                except Exception as e:
                    if retry == max_retries - 1:
                        raise
            
            # Add delay between batches to avoid rate limits
            time.sleep(1)  # 1 second delay between batches
            
        return all_embeddings
        
    except requests.exceptions.HTTPError as http_err:
        logger.error(f"Mistral Embed API HTTP error: {http_err}")
        logger.error(f"Response content: {http_err.response.text if hasattr(http_err, 'response') else 'No response content'}")
        st.error(f"Mistral Embed API HTTP error: {http_err}")
        return []
    except Exception as e:
        logger.error(f"Mistral Embed API error: {str(e)}")
        st.error(f"Mistral Embed API error: {str(e)}")
        return []

def store_embeddings(chunks, vectorizer_type):
    """Store document chunks and their embeddings with vectorizer persistence."""
    try:
        if not chunks:
            raise ValueError("No chunks provided for embedding")
        
        # Clear existing embeddings
        cur.execute("DELETE FROM embeddings")
        conn.commit()
        
        # Filter out empty chunks
        valid_chunks = [chunk for chunk in chunks if chunk and chunk.strip()]
        if not valid_chunks:
            raise ValueError("No valid text content found in chunks")
        
        if vectorizer_type == "TF-IDF":
            # Combine all chunks to fit vectorizer
            all_text = " ".join(valid_chunks)
            
            # Fit the vectorizer and store it
            vectorizer.fit([all_text])
            vectorizer_binary = pickle.dumps(vectorizer)
            
            # Store or update the vectorizer in the database
            cur.execute("""
                INSERT INTO model_state (id, vectorizer)
                VALUES (1, %s)
                ON CONFLICT (id) 
                DO UPDATE SET vectorizer = EXCLUDED.vectorizer, last_updated = CURRENT_TIMESTAMP
            """, (vectorizer_binary,))
            
            successful_inserts = 0
            failed_inserts = 0
            
            for chunk in valid_chunks:
                if not chunk.strip():
                    continue
                    
                embedding = vectorizer.transform([chunk]).toarray()[0].tolist()
                try:
                    cur.execute(
                        "INSERT INTO embeddings (chunk, embedding) VALUES (%s, %s)",
                        (chunk, embedding)
                    )
                    successful_inserts += 1
                except Exception as e:
                    failed_inserts += 1
                    st.error(f"Failed to insert chunk: {str(e)}")
            
            conn.commit()
            st.success(f"Successfully processed {successful_inserts} chunks")
            if failed_inserts > 0:
                st.warning(f"Failed to process {failed_inserts} chunks")
            
            return True
        else:
            # Use Mistral-Embed API to get embeddings
            embeddings = call_mistral_embed_api(valid_chunks)
            if not embeddings:
                st.error("Failed to get embeddings from Mistral-Embed API")
                return False
            
            successful_inserts = 0
            failed_inserts = 0
            
            for chunk, embedding in zip(valid_chunks, embeddings):
                if not chunk.strip():
                    continue
                try:
                    cur.execute(
                        "INSERT INTO embeddings (chunk, embedding) VALUES (%s, %s)",
                        (chunk, embedding)
                    )
                    successful_inserts += 1
                except Exception as e:
                    failed_inserts += 1
                    st.error(f"Failed to insert chunk: {str(e)}")
            
            conn.commit()
            st.success(f"Successfully processed {successful_inserts} chunks")
            if failed_inserts > 0:
                st.warning(f"Failed to process {failed_inserts} chunks")
            
            # Store embeddings as numpy array in session state
            if embeddings:
                st.session_state.current_embeddings = np.array(embeddings)
            
            return True
        
    except Exception as e:
        st.error(f"Error in store_embeddings: {str(e)}")
        conn.rollback()
        return False

def retrieve_relevant_chunks(query, vectorizer_type, top_k=5):
    """Retrieve relevant chunks using cosine similarity."""
    try:
        # Verify embeddings exist
        cur.execute("SELECT COUNT(*) FROM embeddings")
        count = cur.fetchone()[0]
        if count == 0:
            st.error("No embeddings found. Please process documents first.")
            return []
        
        # For empty query, return most recent chunks
        if not query.strip():
            cur.execute("SELECT chunk FROM embeddings ORDER BY id DESC LIMIT %s", (top_k,))
            chunks = cur.fetchall()
            return [{"chunk": chunk[0], "similarity": 1.0} for chunk in chunks]
        
        if vectorizer_type == "TF-IDF":
            if not hasattr(vectorizer, 'vocabulary_'):
                st.error("Vectorizer not fitted! Please initialize document embeddings first.")
                return []
            
            # Generate query embedding
            query_embedding = vectorizer.transform([query]).toarray()[0].tolist()
        else:
            # Use Mistral-Embed API to get query embedding
            query_embeddings = call_mistral_embed_api([query])
            if not query_embeddings:
                return []
            query_embedding = query_embeddings[0]
        
        # Fetch all embeddings
        cur.execute("SELECT chunk, embedding FROM embeddings")
        rows = cur.fetchall()
        
        if not rows:
            st.warning("No embeddings found in database. Please initialize document embeddings first.")
            return []
        
        # Split chunks and embeddings
        chunks = []
        embeddings = []
        for row in rows:
            if row[0] and row[1]:
                chunks.append(row[0])
                # Ensure embedding is a list of floats with consistent length
                embedding = [float(x) for x in row[1]]
                embeddings.append(embedding)
        
        if not chunks:
            st.warning("No valid chunks found in database")
            return []
        
        # Ensure all embeddings have the same length
        embedding_length = len(embeddings[0])
        embeddings = [emb for emb in embeddings if len(emb) == embedding_length]
        
        # Convert to numpy array with explicit dtype
        embeddings_array = np.array(embeddings, dtype=np.float32)
        query_embedding_array = np.array(query_embedding, dtype=np.float32).reshape(1, -1)
        
        # Calculate similarities
        similarities = cosine_similarity(query_embedding_array, embeddings_array)[0]
        
        # Get top results
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        results = [{"chunk": chunks[i], "similarity": float(similarities[i])} for i in top_indices]
        
        return results
        
    except Exception as e:
        st.error(f"Error in retrieve_relevant_chunks: {str(e)}")
        logger.error(f"Error retrieving chunks: {str(e)}")
        return []

def call_mistral_api(prompt, model, temperature, max_tokens):
    """Call Mistral API with error handling."""
    try:
        headers = {
            "Authorization": f"Bearer {MISTRAL_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        response = requests.post(MISTRAL_API_ENDPOINT, headers=headers, json=data)
        response.raise_for_status()
        return response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
    except Exception as e:
        logger.error(f"Mistral API error: {str(e)}")
        return f"Error calling Mistral API: {str(e)}"

def diagnose_document_processing(vectorizer_type):
    """Diagnostic function to check document processing pipeline."""
    with st.expander("Diagnostics Results", expanded=True):
        st.write("Running diagnostics...")
        
        # Check S3 connection and folder contents
        try:
            response = s3_client.list_objects_v2(
                Bucket=S3_BUCKET_NAME,
                Prefix=S3_FOLDER_NAME
            )
            st.write("✅ S3 connection successful")
            folder_files = [
                obj['Key'] for obj in response.get('Contents', [])
                if obj['Key'].startswith(S3_FOLDER_NAME)
            ]
            if folder_files:
                st.write(f"Found {len(folder_files)} objects in folder {S3_FOLDER_NAME}")
            else:
                st.warning(f"No objects found in folder {S3_FOLDER_NAME}")
        except Exception as e:
            st.error(f"S3 connection failed: {str(e)}")
        
        # Check vectorizer/embedding state
        try:
            if vectorizer_type == "TF-IDF":
                if hasattr(vectorizer, 'vocabulary_'):
                    st.write("✅ TF-IDF Vectorizer is fitted")
                    st.write(f"Vocabulary size: {len(vectorizer.vocabulary_)}")
                else:
                    st.warning("❌ TF-IDF Vectorizer is not fitted")
            else:
                st.write("✅ Using Mistral Embed API for embeddings")
        except Exception as e:
            st.error(f"Error checking vectorizer state: {str(e)}")
        
        # Check database state
        try:
            cur.execute("SELECT COUNT(*) FROM embeddings")
            count = cur.fetchone()[0]
            st.write(f"✅ Database connected, found {count} embeddings")
            
            if count > 0:
                cur.execute("SELECT chunk FROM embeddings LIMIT 1")
                sample_chunk = cur.fetchone()[0]
                st.write("Sample chunk preview:")
                st.write(sample_chunk[:200] + "...")
        except Exception as e:
            st.error(f"Database check failed: {str(e)}")

def init_session_state():
    """Initialize session state variables"""
    if 'history' not in st.session_state:
        st.session_state.history = []
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if 'cache' not in st.session_state:
        st.session_state.cache = {}
    if 'documents_processed' not in st.session_state:
        st.session_state.documents_processed = False
    if 'processed_files' not in st.session_state:
        st.session_state.processed_files = []
    if 'vectorizer' not in st.session_state:
        st.session_state.vectorizer = None

@st.cache_data(ttl=3600)
def get_cached_response(prompt, context):
    """Cache responses to avoid repeated API calls"""
    cache_key = f"{prompt}_{hash(context)}"
    return st.session_state.cache.get(cache_key)

def store_cached_response(prompt, context, response):
    """Store response in cache"""
    cache_key = f"{prompt}_{hash(context)}"
    st.session_state.cache[cache_key] = response

def visualize_embeddings(embeddings, labels=None):
    """Create t-SNE visualization of embeddings"""
    try:
        from sklearn.manifold import TSNE
        
        # Convert embeddings to numpy array if it's a list
        embeddings_array = np.array(embeddings) if isinstance(embeddings, list) else embeddings
        
        # Check if we have valid embeddings
        if embeddings_array.size == 0:
            st.warning("No embeddings available for visualization")
            return None
            
        # Apply t-SNE
        tsne = TSNE(n_components=2, random_state=42)
        embeddings_2d = tsne.fit_transform(embeddings_array)
        
        # Create visualization
        df = pd.DataFrame(embeddings_2d, columns=['x', 'y'])
        if labels:
            df['label'] = labels[:len(df)] if len(labels) > len(df) else labels
        
        fig = px.scatter(
            df, 
            x='x', 
            y='y', 
            text='label' if labels else None,
            title='Document Embeddings Visualization',
            labels={'x': 'TSNE-1', 'y': 'TSNE-2'}
        )
        return fig
    except Exception as e:
        logger.error(f"Visualization error: {str(e)}")
        st.warning("Could not create embedding visualization")
        return None

def file_uploader_ui():
    """Handle file uploads with progress tracking and validation"""
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    ALLOWED_TYPES = {
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }

    uploaded_files = st.file_uploader(
        "Upload documents", 
        accept_multiple_files=True,
        type=list(ALLOWED_TYPES.keys()),
        help="Upload PDF, TXT, or DOC files (max 10MB each)"
    )
    
    if uploaded_files:
        upload_status = st.empty()
        progress_bar = st.progress(0)
        
        successful_uploads = 0
        failed_uploads = 0
        total_files = len(uploaded_files)
        
        for idx, file in enumerate(uploaded_files):
            try:
                # Validate file
                if not file:
                    raise ValueError("Invalid file object")
                
                logger.info(f"Processing file: {file.name}, Size: {file.size} bytes, Type: {file.type}")
                
                if file.size > MAX_FILE_SIZE:
                    raise ValueError(f"File size ({file.size / 1024 / 1024:.2f}MB) exceeds 10MB limit")
                
                if file.size == 0:
                    raise ValueError("File is empty")
                
                file_extension = file.name.split('.')[-1].lower()
                if file_extension not in ALLOWED_TYPES:
                    raise ValueError(f"Unsupported file type: {file_extension}")
                
                progress = (idx + 1) / total_files
                progress_bar.progress(progress)
                upload_status.text(f"Uploading {file.name}... ({idx + 1}/{total_files})")
                
                file.seek(0)
                
                # Upload directly to bucket root
                s3_key = f"{S3_FOLDER_NAME}{file.name}"
                
                try:
                    s3_client.upload_fileobj(
                        file,
                        S3_BUCKET_NAME,
                        s3_key,
                        ExtraArgs={
                            'ContentType': ALLOWED_TYPES.get(file_extension, 'application/octet-stream')
                        }
                    )
                except Exception as s3_error:
                    raise ValueError(f"S3 upload failed: {str(s3_error)}")
                
                successful_uploads += 1
                st.success(f"✅ Successfully uploaded {file.name}")
                
            except Exception as e:
                failed_uploads += 1
                error_msg = f"❌ Failed to upload {file.name}: {str(e)}"
                st.error(error_msg)
                logger.error(f"Upload error: {error_msg}")
                logger.exception("Detailed error traceback:")
                continue
        
        progress_bar.empty()
        upload_status.empty()
        
        if successful_uploads > 0:
            st.success(f"🎉 Successfully uploaded {successful_uploads} files")
        if failed_uploads > 0:
            st.warning(f"⚠️ Failed to upload {failed_uploads} files")
            st.info("Please check file size, format, and try again")
        
        return successful_uploads > 0

def get_task_prompt(context, task_type, **kwargs):
    """Generate appropriate prompt based on task type."""
    prompts = {
        "summarization": {
            "concise": f"Provide a concise summary of the following text in 2-3 paragraphs:\n\n{context}",
            "detailed": f"Provide a detailed summary of the following text, including key points and main ideas:\n\n{context}",
            "bullet_points": f"Summarize the following text in bullet points, highlighting the main ideas:\n\n{context}",
            "executive": f"Provide an executive summary of the following text, focusing on key findings and conclusions:\n\n{context}"
        },
        "classification": f"Classify the following text into appropriate categories. Consider tone, topic, and intent:\n\n{context}",
        "prediction": f"Based on the following text, predict {kwargs.get('prediction_target', 'future implications')}:\n\n{context}"
    }
    
    if task_type == "summarization":
        style = kwargs.get("style", "concise")
        return prompts["summarization"][style]
    return prompts.get(task_type, prompts["summarization"]["concise"])

def main():
    init_session_state()
    initialize_database()
    
    st.set_page_config(page_title="AI Document Analysis", layout="wide")
    
    # Title with center alignment using markdown
    st.markdown("<h1 style='text-align: center;'>AI: Summarize, Classify, Predict</h1>", unsafe_allow_html=True)
    
    # Left-aligned description with bullet points on separate lines
    st.markdown(
        "A Retrieval-Augmented Generation (RAG) application using Streamlit and AWS allows users to summarize key information from uploaded or entered text, classify documents based on their content and context, and predict future outcomes or trends"
    )
    
    # Initialize session state for document processing
    if 'documents_processed' not in st.session_state:
        st.session_state.documents_processed = False
    if 'processed_files' not in st.session_state:
        st.session_state.processed_files = []
    
    # Three column layout
    col1, col2, col3 = st.columns([1, 1, 1])
    
    # Store UI state variables
    state = {
        'task_type': None,
        'style': None,
        'prediction_target': None,
        'model': None,
        'temperature': None,
        'max_tokens': None,
        'input_source': None,
        'vectorizer_type': None,
        'chunking_strategy': None,
        'selected_s3_files': None,
        'user_text': None
    }
    
    with col1:
        # Analysis Type column
        st.markdown("### 📊 1. Analysis Type")
        state['task_type'] = st.selectbox(
            "Select Analysis Type",
            ["summarization", "classification", "prediction"],
            help="Choose the type of analysis you want to perform"
        )
        
        if state['task_type'] == "summarization":
            state['style'] = st.selectbox(
                "Summary Style", 
                ["concise", "detailed", "bullet_points", "executive"],
                help="Choose how you want your summary formatted"
            )
        elif state['task_type'] == "classification":
            st.info("Will classify document content by topic, tone, and intent")
        elif state['task_type'] == "prediction":
            state['prediction_target'] = st.text_input(
                "Prediction Target", 
                "future implications",
                help="What aspects would you like to predict?"
            )
        
        state['model'] = st.selectbox(
            "Model", 
            ["mistral-small-latest", "open-mistral-7b"],
            help="Select the AI model to use"
        )
        state['temperature'] = st.slider("Creativity", 0.0, 1.0, 0.3)
        state['max_tokens'] = st.slider("Max Length", 100, 2000, 500)
    
    with col2:
        # Input Sources column
        st.markdown("### 📁 2. Input Sources")
        state['input_source'] = st.radio(
            "Select Input Source",
            ["S3 Documents", "Upload Documents", "Both", "Direct Text Input"],
            help="Choose where to get the documents from"
        )
        
        # Common options for both S3 and Upload
        state['vectorizer_type'] = st.selectbox("Embedding Method", ["Mistral-Embed", "TF-IDF"])
        state['chunking_strategy'] = st.selectbox("Chunking Strategy", ["Sentence-Based", "Fixed-Size"])
        
        # Add optional context for document-based analysis
        if state['input_source'] in ["S3 Documents", "Upload Documents", "Both"]:
            state['additional_context'] = st.text_area(
                "Additional Context (Optional)",
                help="Add any additional context or instructions to consider during analysis",
                height=100
            )
        
        # Rest of the input source handling
        if state['input_source'] in ["Upload Documents", "Both"]:
            uploaded = file_uploader_ui()
        
        if state['input_source'] in ["S3 Documents", "Both"]:
            s3_files = get_s3_files()
            if s3_files:
                state['selected_s3_files'] = st.multiselect(
                    "Select S3 Documents",
                    s3_files,
                    help="Choose documents from S3 bucket"
                )
                if not state['selected_s3_files']:
                    st.info("Please select at least one document")
            else:
                st.warning(f"No documents found in folder {S3_FOLDER_NAME}")
                st.info("Upload documents first or choose a different input source")
        
        if state['input_source'] == "Direct Text Input":
            state['user_text'] = st.text_area(
                "Enter your text",
                height=200,
                help="Paste or type the text you want to analyze"
            )
    
    with col3:
        # Process column
        st.markdown("### ⚙️ 3. Process")
        
        # Process button for document sources
        if state['input_source'] in ["S3 Documents", "Upload Documents", "Both"]:
            if st.button("Process Documents", use_container_width=True):
                with st.spinner("Processing documents..."):
                    process_documents(
                        state['input_source'], 
                        state['selected_s3_files'],
                        state['vectorizer_type'],
                        state['chunking_strategy'],
                        state.get('additional_context', '')  # Pass additional context
                    )
        
        # Generate button
        generate_button_label = f"Generate {state['task_type'].title()}"
        if st.button(generate_button_label, use_container_width=True):
            with st.spinner(f"Generating {state['task_type']}..."):
                if state['input_source'] == "Direct Text Input":
                    if state['user_text']:
                        result = generate_from_text(
                            state['user_text'],
                            state['task_type'],
                            state['model'],
                            state['temperature'],
                            state['max_tokens']
                        )
                        display_results(result, state['task_type'])
                    else:
                        st.error("Please enter some text to analyze")
                else:
                    generate_from_processed_documents(
                        state['task_type'],
                        state['model'],
                        state['temperature'],
                        state['max_tokens'],
                        state['style'] if state['task_type'] == "summarization" else None
                    )
        
        # Display processing status and results
        if st.session_state.get('documents_processed'):
            st.success("✅ Documents processed and ready for analysis")
        
        # Show history in expander
        if st.session_state.history:
            with st.expander("📜 Analysis History"):
                for item in reversed(st.session_state.history[-5: ]):
                    st.markdown(f"**{item['timestamp']} - {item['task_type']}**")
                    st.write(item['result'])
                    st.markdown("---")

# ...rest of the existing code...

def get_s3_files():
    """Get list of files from S3 bucket's App2 folder"""
    try:
        response = s3_client.list_objects_v2(
            Bucket=S3_BUCKET_NAME,
            Prefix=S3_FOLDER_NAME
        )
        files = []
        for obj in response.get('Contents', []):
            if obj['Key'].startswith(S3_FOLDER_NAME) and obj['Key'].lower().endswith(('.pdf', '.txt', '.doc', '.docx')):
                files.append(obj['Key'])
        logger.info(f"Found {len(files)} files in {S3_FOLDER_NAME}")
        return files
    except Exception as e:
        logger.error(f"Error accessing S3: {str(e)}")
        return []

def process_documents(input_source, s3_files=None, vectorizer_type="Mistral-Embed", chunking_strategy="Sentence-Based", additional_context=None):
    """Process documents from various sources"""
    try:
        if input_source in ["S3 Documents", "Both"]:
            if not s3_files:
                st.error("Please select at least one document to process")
                return False
                
            # Verify all files are from correct folder
            if not all(file.startswith(S3_FOLDER_NAME) for file in s3_files):
                st.error(f"Some selected files are not from the {S3_FOLDER_NAME} folder")
                return False
        
        all_chunks = []
        processed_files = []
        
        # Add additional context as first chunk if provided
        if additional_context and additional_context.strip():
            all_chunks.append(additional_context.strip())
            logger.info("Added additional context to processing")
        
        # Process S3 documents
        if input_source in ["S3 Documents", "Both"] and s3_files:
            for file_key in s3_files:
                with st.spinner(f"Processing {file_key}..."):
                    document_content = load_document_from_s3(S3_BUCKET_NAME, file_key)
                    if document_content:
                        chunks = chunk_document(document_content, chunking_strategy)
                        if chunks:
                            all_chunks.extend(chunks)
                            processed_files.append(file_key)
                            logger.info(f"Successfully chunked {file_key} into {len(chunks)} chunks")
                        else:
                            st.warning(f"No valid content extracted from {file_key}")
        
        if not all_chunks:
            st.error("No valid content was extracted from the documents")
            return False
            
        # Store embeddings
        with st.spinner("Generating embeddings..."):
            if store_embeddings(all_chunks, vectorizer_type):
                st.session_state.processed_files = processed_files
                st.session_state.documents_processed = True
                st.success(f"✅ Successfully processed {len(processed_files)} documents with additional context")
                return True
            
        return False
            
    except Exception as e:
        st.error(f"Error processing documents: {str(e)}")
        logger.error(f"Document processing error: {str(e)}")
        return False

def generate_from_processed_documents(task_type, model, temperature, max_tokens, style=None):
    """Generate analysis from processed documents"""
    try:
        if not st.session_state.documents_processed:
            st.error("Please process documents first")
            return
            
        chunks = retrieve_relevant_chunks("", "Mistral-Embed", top_k=10)
        if not chunks:
            st.error("No processed content available. Please ensure documents are properly processed.")
            return
            
        context = " ".join(chunk["chunk"] for chunk in chunks)
        prompt = get_task_prompt(context, task_type, style=style)
        result = call_mistral_api(prompt, model, temperature, max_tokens)
        if result:
            display_results(result, task_type)
        else:
            st.error("Failed to generate analysis")
            
    except Exception as e:
        st.error(f"Error generating analysis: {str(e)}")

def generate_from_text(text, task_type, model, temperature, max_tokens):
    """Generate analysis from direct text input"""
    prompt = get_task_prompt(text, task_type)
    return call_mistral_api(prompt, model, temperature, max_tokens)

def display_results(result, task_type):
    """Display analysis results"""
    st.markdown("### Results")
    st.write(result)
    
    # Add to history
    st.session_state.history.append({
        'timestamp': datetime.now().isoformat(),
        'task_type': task_type,
        'result': result
    })
    
    # Note: Embedding visualization removed as requested

if __name__ == "__main__":
    main()