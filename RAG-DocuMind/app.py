import streamlit as st
import boto3
import psycopg2
import requests
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import PyPDF2
import io
import os
from dotenv import load_dotenv
import logging
import pickle
import time  # Add this import
import re  # Add this import
import plotly.express as px
from datetime import datetime
import uuid
import json
from streamlit import session_state
from docx import Document  # Add this import

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration from environment variables
S3_BUCKET_NAME = os.getenv('S3_BUCKET_NAME')
S3_DOCUMENT_KEY = os.getenv('S3_DOCUMENT_KEY', '')
AWS_ACCESS_KEY = os.getenv('AWS_ACCESS_KEY')
AWS_SECRET_KEY = os.getenv('AWS_SECRET_KEY')
AWS_REGION = os.getenv('AWS_REGION')
POSTGRES_HOST = os.getenv('POSTGRES_HOST')
POSTGRES_PORT = os.getenv('POSTGRES_PORT')
POSTGRES_DB = os.getenv('POSTGRES_DB')
POSTGRES_USER = os.getenv('POSTGRES_USER')
POSTGRES_PASSWORD = os.getenv('POSTGRES_PASSWORD')
MISTRAL_API_KEY = os.getenv('MISTRAL_API_KEY')
MISTRAL_API_ENDPOINT = "https://api.mistral.ai/v1/chat/completions"
MISTRAL_EMBED_API_ENDPOINT = "https://api.mistral.ai/v1/embeddings"

# Initialize S3 Client
s3_client = boto3.client(
    's3',
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region_name=AWS_REGION
)

# Initialize PostgreSQL Connection
def get_db_connection():
    try:
        conn = psycopg2.connect(
            host=POSTGRES_HOST,
            port=POSTGRES_PORT,
            dbname=POSTGRES_DB,
            user=POSTGRES_USER,
            password=POSTGRES_PASSWORD
        )
        return conn
    except Exception as e:
        logger.error(f"Database connection error: {str(e)}")
        raise

conn = get_db_connection()
cur = conn.cursor()

# Initialize database tables
def initialize_database():
    """Initialize database tables with proper error handling."""
    try:
        # Create tables with IF NOT EXISTS clause
        cur.execute('''
            CREATE TABLE IF NOT EXISTS embeddings (
                id SERIAL PRIMARY KEY,
                chunk TEXT,
                embedding FLOAT[]
            )
        ''')
        
        cur.execute('''
            CREATE TABLE IF NOT EXISTS model_state (
                id INTEGER PRIMARY KEY,
                vectorizer BYTEA,
                last_updated TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Ensure the tables were created successfully
        cur.execute("""
            SELECT EXISTS (
                SELECT FROM information_schema.tables 
                WHERE table_name = 'model_state'
            )
        """)
        model_state_exists = cur.fetchone()[0]
        
        cur.execute("""
            SELECT EXISTS (
                SELECT FROM information_schema.tables 
                WHERE table_name = 'embeddings'
            )
        """)
        embeddings_exists = cur.fetchone()[0]
        
        if not (model_state_exists and embeddings_exists):
            raise Exception("Failed to create required tables")
            
        conn.commit()
        logger.info("Database tables initialized successfully")
        
    except Exception as e:
        logger.error(f"Database initialization error: {str(e)}")
        conn.rollback()
        raise

# Modified vectorizer initialization with persistence
@st.cache_resource
def initialize_vectorizer(vectorizer_type):
    try:
        if vectorizer_type == "TF-IDF":
            # Try to load fitted vectorizer from database
            cur.execute("SELECT vectorizer FROM model_state WHERE id = 1")
            result = cur.fetchone()
            if result and result[0]:
                vectorizer = pickle.loads(result[0])
                st.success("Loaded previously fitted TF-IDF vectorizer")
                return vectorizer, True
            else:
                vectorizer = TfidfVectorizer(
                    max_features=100,
                    stop_words='english',
                    lowercase=True,
                    dtype=np.float32
                )
                return vectorizer, False
        else:
            # For Mistral-Embed, no need to load from database
            return None, False
    except Exception as e:
        st.error(f"Error initializing vectorizer: {str(e)}")
        if vectorizer_type == "TF-IDF":
            vectorizer = TfidfVectorizer(
                max_features=100,
                stop_words='english',
                lowercase=True,
                dtype=np.float32
            )
            return vectorizer, False
        else:
            return None, False

def extract_text_from_pdf(pdf_content):
    """Extract text from PDF binary content."""
    try:
        pdf_file = io.BytesIO(pdf_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return None

def extract_text_from_docx(docx_content):
    """Extract text from DOCX binary content."""
    try:
        doc = Document(io.BytesIO(docx_content))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return None

def load_document_from_s3(bucket_name, document_key):
    """Load document from S3 and handle multiple document types."""
    try:
        response = s3_client.get_object(Bucket=bucket_name, Key=document_key)
        document_content = response['Body'].read()
        
        # Log document details for debugging
        logger.info(f"Loading document: {document_key}, Size: {len(document_content)} bytes")
        
        file_extension = document_key.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return extract_text_from_pdf(document_content)
        elif file_extension in ['docx', 'doc']:
            return extract_text_from_docx(document_content)
        else:
            try:
                return document_content.decode('utf-8')
            except UnicodeDecodeError:
                return document_content.decode('latin-1')
    except Exception as e:
        logger.error(f"Error loading document from S3: {str(e)}")
        st.error(f"Error loading document {document_key}: {str(e)}")
        return None

def chunk_document_fixed_size(document, chunk_size=512, overlap=50):
    """Fixed-size chunking with sliding window."""
    if not document:
        return []
        
    chunks = []
    words = document.split()
    
    chunk_size_words = chunk_size // 4  # Approximate words per chunk
    for i in range(0, len(words), chunk_size_words - overlap):
        chunk = ' '.join(words[i:i + chunk_size_words])
        if len(chunk) > 8000:
            chunk = chunk[:8000]
        chunks.append(chunk)
    return chunks

def chunk_document_sentence(document, max_chunk_size=8000):
    """Sentence-based chunking preserving semantic meaning."""
    if not document:
        return []
    
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', document)
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        sentence_size = len(sentence)
        
        # If single sentence exceeds max size, split it into smaller chunks
        if sentence_size > max_chunk_size:
            if current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_size = 0
            
            # Split long sentence
            words = sentence.split()
            temp_chunk = []
            temp_size = 0
            
            for word in words:
                if temp_size + len(word) + 1 > max_chunk_size:
                    chunks.append(' '.join(temp_chunk))
                    temp_chunk = [word]
                    temp_size = len(word)
                else:
                    temp_chunk.append(word)
                    temp_size += len(word) + 1
            
            if temp_chunk:
                chunks.append(' '.join(temp_chunk))
            
        # If adding sentence exceeds limit, create new chunk
        elif current_size + sentence_size + 1 > max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [sentence]
            current_size = sentence_size
        else:
            current_chunk.append(sentence)
            current_size += sentence_size + 1
    
    # Add the last chunk
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def chunk_document(document, chunking_strategy, chunk_size=512, overlap=50):
    """Chunk document using selected strategy."""
    if chunking_strategy == "Fixed-Size":
        return chunk_document_fixed_size(document, chunk_size, overlap)
    else:
        return chunk_document_sentence(document)

def call_mistral_embed_api(texts):
    """Call Mistral Embed API to get embeddings."""
    try:
        # Ensure texts are non-empty and properly formatted
        texts = [text[:8000] for text in texts if text.strip()]  # Limit text length to 8000 chars
        if not texts:
            raise ValueError("No valid texts provided for embedding.")
        
        # Process in smaller batches with delay between calls
        batch_size = 5  # Reduced batch size
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            
            headers = {
                "Authorization": f"Bearer {MISTRAL_API_KEY}",
                "Content-Type": "application/json"
            }
            data = {
                "input": batch[0] if len(batch) == 1 else batch,
                "model": "mistral-embed"
            }
            
            # Log the request payload for debugging
            logger.info(f"Request payload for batch {i//batch_size + 1}: {data}")
            
            max_retries = 3
            retry_delay = 5  # seconds
            
            for retry in range(max_retries):
                try:
                    response = requests.post(MISTRAL_EMBED_API_ENDPOINT, headers=headers, json=data)
                    
                    if response.status_code == 429:  # Too Many Requests
                        wait_time = int(response.headers.get('Retry-After', retry_delay * (retry + 1)))
                        logger.warning(f"Rate limit hit, waiting {wait_time} seconds...")
                        time.sleep(wait_time)
                        continue
                    
                    response.raise_for_status()
                    embeddings = response.json().get("data", [])
                    all_embeddings.extend([e.get("embedding") for e in embeddings])
                    
                    # Log success for this batch
                    logger.info(f"Successfully processed batch {i//batch_size + 1}")
                    break  # Success, exit retry loop
                    
                except requests.exceptions.HTTPError as http_err:
                    if response.status_code != 429 or retry == max_retries - 1:
                        raise
                
                except Exception as e:
                    if retry == max_retries - 1:
                        raise
            
            # Add delay between batches to avoid rate limits
            time.sleep(1)  # 1 second delay between batches
            
        return all_embeddings
        
    except requests.exceptions.HTTPError as http_err:
        logger.error(f"Mistral Embed API HTTP error: {http_err}")
        logger.error(f"Response content: {http_err.response.text if hasattr(http_err, 'response') else 'No response content'}")
        st.error(f"Mistral Embed API HTTP error: {http_err}")
        return []
    except Exception as e:
        logger.error(f"Mistral Embed API error: {str(e)}")
        st.error(f"Mistral Embed API error: {str(e)}")
        return []

def store_embeddings(chunks, vectorizer_type):
    """Store document chunks and their embeddings with vectorizer persistence."""
    try:
        # Clear existing embeddings
        cur.execute("DELETE FROM embeddings")
        
        if vectorizer_type == "TF-IDF":
            # Combine all chunks to fit vectorizer
            all_text = " ".join(chunks)
            
            # Fit the vectorizer and store it
            vectorizer.fit([all_text])
            vectorizer_binary = pickle.dumps(vectorizer)
            
            # Store or update the vectorizer in the database
            cur.execute("""
                INSERT INTO model_state (id, vectorizer)
                VALUES (1, %s)
                ON CONFLICT (id) 
                DO UPDATE SET vectorizer = EXCLUDED.vectorizer, last_updated = CURRENT_TIMESTAMP
            """, (vectorizer_binary,))
            
            successful_inserts = 0
            failed_inserts = 0
            
            for chunk in chunks:
                if not chunk.strip():
                    continue
                    
                embedding = vectorizer.transform([chunk]).toarray()[0].tolist()
                try:
                    cur.execute(
                        "INSERT INTO embeddings (chunk, embedding) VALUES (%s, %s)",
                        (chunk, embedding)
                    )
                    successful_inserts += 1
                except Exception as e:
                    failed_inserts += 1
                    st.error(f"Failed to insert chunk: {str(e)}")
            
            conn.commit()
            st.success(f"Successfully processed {successful_inserts} chunks")
            if failed_inserts > 0:
                st.warning(f"Failed to process {failed_inserts} chunks")
            
            return True
        else:
            # Use Mistral-Embed API to get embeddings
            embeddings = call_mistral_embed_api(chunks)
            if not embeddings:
                st.error("Failed to get embeddings from Mistral-Embed API")
                return False
            
            successful_inserts = 0
            failed_inserts = 0
            
            for chunk, embedding in zip(chunks, embeddings):
                if not chunk.strip():
                    continue
                try:
                    cur.execute(
                        "INSERT INTO embeddings (chunk, embedding) VALUES (%s, %s)",
                        (chunk, embedding)
                    )
                    successful_inserts += 1
                except Exception as e:
                    failed_inserts += 1
                    st.error(f"Failed to insert chunk: {str(e)}")
            
            conn.commit()
            st.success(f"Successfully processed {successful_inserts} chunks")
            if failed_inserts > 0:
                st.warning(f"Failed to process {failed_inserts} chunks")
            
            return True
        
    except Exception as e:
        st.error(f"Error in store_embeddings: {str(e)}")
        conn.rollback()
        return False

def retrieve_relevant_chunks(query, vectorizer_type, top_k=5):
    """Retrieve relevant chunks using cosine similarity."""
    try:
        if vectorizer_type == "TF-IDF":
            if not hasattr(vectorizer, 'vocabulary_'):
                st.error("Vectorizer not fitted! Please initialize document embeddings first.")
                return []
            
            # Generate query embedding
            query_embedding = vectorizer.transform([query]).toarray()[0].tolist()
        else:
            # Use Mistral-Embed API to get query embedding
            query_embeddings = call_mistral_embed_api([query])
            if not query_embeddings:
                return []
            query_embedding = query_embeddings[0]
        
        # Fetch all embeddings
        cur.execute("SELECT chunk, embedding FROM embeddings")
        rows = cur.fetchall()
        
        if not rows:
            st.warning("No embeddings found in database. Please initialize document embeddings first.")
            return []
        
        # Split chunks and embeddings
        chunks = []
        embeddings = []
        for row in rows:
            if row[0] and row[1]:
                chunks.append(row[0])
                # Ensure embedding is a list of floats with consistent length
                embedding = [float(x) for x in row[1]]
                embeddings.append(embedding)
        
        if not chunks:
            st.warning("No valid chunks found in database")
            return []
        
        # Ensure all embeddings have the same length
        embedding_length = len(embeddings[0])
        embeddings = [emb for emb in embeddings if len(emb) == embedding_length]
        
        # Convert to numpy array with explicit dtype
        embeddings_array = np.array(embeddings, dtype=np.float32)
        query_embedding_array = np.array(query_embedding, dtype=np.float32).reshape(1, -1)
        
        # Calculate similarities
        similarities = cosine_similarity(query_embedding_array, embeddings_array)[0]
        
        # Get top results
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        results = [{"chunk": chunks[i], "similarity": float(similarities[i])} for i in top_indices]
        
        return results
        
    except Exception as e:
        st.error(f"Error in retrieve_relevant_chunks: {str(e)}")
        logger.error(f"Error retrieving chunks: {str(e)}")
        return []

def call_mistral_api(prompt, model, temperature, max_tokens):
    """Call Mistral API with error handling."""
    try:
        headers = {
            "Authorization": f"Bearer {MISTRAL_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        response = requests.post(MISTRAL_API_ENDPOINT, headers=headers, json=data)
        response.raise_for_status()
        return response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
    except Exception as e:
        logger.error(f"Mistral API error: {str(e)}")
        return f"Error calling Mistral API: {str(e)}"

def diagnose_document_processing(vectorizer_type):
    """Diagnostic function to check document processing pipeline."""
    with st.expander("Diagnostics Results", expanded=True):
        st.write("Running diagnostics...")
        
        # Check S3 connection
        try:
            response = s3_client.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=S3_DOCUMENT_KEY)
            st.write("✅ S3 connection successful")
            if 'Contents' in response:
                st.write(f"Found {len(response['Contents'])} objects with prefix {S3_DOCUMENT_KEY}")
            else:
                st.warning(f"No objects found with prefix {S3_DOCUMENT_KEY}")
        except Exception as e:
            st.error(f"S3 connection failed: {str(e)}")
        
        # Check vectorizer/embedding state
        try:
            if vectorizer_type == "TF-IDF":
                if hasattr(vectorizer, 'vocabulary_'):
                    st.write("✅ TF-IDF Vectorizer is fitted")
                    st.write(f"Vocabulary size: {len(vectorizer.vocabulary_)}")
                else:
                    st.warning("❌ TF-IDF Vectorizer is not fitted")
            else:
                st.write("✅ Using Mistral Embed API for embeddings")
        except Exception as e:
            st.error(f"Error checking vectorizer state: {str(e)}")
        
        # Check database state
        try:
            cur.execute("SELECT COUNT(*) FROM embeddings")
            count = cur.fetchone()[0]
            st.write(f"✅ Database connected, found {count} embeddings")
            
            if count > 0:
                cur.execute("SELECT chunk FROM embeddings LIMIT 1")
                sample_chunk = cur.fetchone()[0]
                st.write("Sample chunk preview:")
                st.write(sample_chunk[:200] + "...")
        except Exception as e:
            st.error(f"Database check failed: {str(e)}")

def init_session_state():
    if 'history' not in st.session_state:
        st.session_state.history = []
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if 'cache' not in st.session_state:
        st.session_state.cache = {}

@st.cache_data(ttl=3600)
def get_cached_response(prompt, context):
    """Cache responses to avoid repeated API calls"""
    cache_key = f"{prompt}_{hash(context)}"
    return st.session_state.cache.get(cache_key)

def store_cached_response(prompt, context, response):
    """Store response in cache"""
    cache_key = f"{prompt}_{hash(context)}"
    st.session_state.cache[cache_key] = response

def visualize_embeddings(embeddings, labels=None):
    """Create t-SNE visualization of embeddings"""
    from sklearn.manifold import TSNE
    tsne = TSNE(n_components=2, random_state=42)
    embeddings_2d = tsne.fit_transform(embeddings)
    
    df = pd.DataFrame(embeddings_2d, columns=['x', 'y'])
    if labels:
        df['label'] = labels
    
    fig = px.scatter(df, x='x', y='y', text=labels if labels else None,
                    title='Document Embeddings Visualization')
    return fig

def file_uploader_ui():
    """Handle file uploads with progress tracking and validation"""
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    ALLOWED_TYPES = {
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }

    uploaded_files = st.file_uploader(
        "Upload documents", 
        accept_multiple_files=True,
        type=list(ALLOWED_TYPES.keys()),
        help="Upload PDF, TXT, or DOC files (max 10MB each)"
    )
    
    if uploaded_files:
        upload_status = st.empty()
        progress_bar = st.progress(0)
        
        successful_uploads = 0
        failed_uploads = 0
        total_files = len(uploaded_files)
        
        for idx, file in enumerate(uploaded_files):
            try:
                # Validate file
                if not file:
                    raise ValueError("Invalid file object")
                
                logger.info(f"Processing file: {file.name}, Size: {file.size} bytes, Type: {file.type}")
                
                if file.size > MAX_FILE_SIZE:
                    raise ValueError(f"File size ({file.size / 1024 / 1024:.2f}MB) exceeds 10MB limit")
                
                if file.size == 0:
                    raise ValueError("File is empty")
                
                file_extension = file.name.split('.')[-1].lower()
                if file_extension not in ALLOWED_TYPES:
                    raise ValueError(f"Unsupported file type: {file_extension}")
                
                progress = (idx + 1) / total_files
                progress_bar.progress(progress)
                upload_status.text(f"Uploading {file.name}... ({idx + 1}/{total_files})")
                
                file.seek(0)
                
                # Upload directly to bucket root
                s3_key = file.name
                
                try:
                    s3_client.upload_fileobj(
                        file,
                        S3_BUCKET_NAME,
                        s3_key,
                        ExtraArgs={
                            'ContentType': ALLOWED_TYPES.get(file_extension, 'application/octet-stream')
                        }
                    )
                except Exception as s3_error:
                    raise ValueError(f"S3 upload failed: {str(s3_error)}")
                
                successful_uploads += 1
                st.success(f"✅ Successfully uploaded {file.name}")
                
            except Exception as e:
                failed_uploads += 1
                error_msg = f"❌ Failed to upload {file.name}: {str(e)}"
                st.error(error_msg)
                logger.error(f"Upload error: {error_msg}")
                logger.exception("Detailed error traceback:")
                continue
        
        progress_bar.empty()
        upload_status.empty()
        
        if successful_uploads > 0:
            st.success(f"🎉 Successfully uploaded {successful_uploads} files")
        if failed_uploads > 0:
            st.warning(f"⚠️ Failed to upload {failed_uploads} files")
            st.info("Please check file size, format, and try again")
        
        return successful_uploads > 0

def main():
    # Initialize session state
    init_session_state()
    
    try:
        # Initialize database tables before any operations
        initialize_database()
    except Exception as e:
        st.error(f"Failed to initialize database: {str(e)}")
        st.warning("Please ensure your database connection is correct and you have proper permissions.")
        return

    # Set page config
    st.set_page_config(
        page_title="DocuMind: Intelligent Document Analysis & Response Platform",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Center-aligned title with smaller size
    st.markdown("<h2 style='text-align: center;'>DocuMind: Intelligent Document Analysis & Response Platform</h2>", unsafe_allow_html=True)
    
    # Application description
    st.markdown("""
        <div style='text-align: center; max-width: 800px; margin: 0 auto; margin-bottom: 20px;'>
            A Retrieval-Augmented Generation (RAG) application leveraging AWS services for document processing and analysis. 
            This application combines the power of large language models with your custom document repository for intelligent document processing, contextual understanding, and  
            automated response generation.
        </div>
    """, unsafe_allow_html=True)
    
    # Key Capabilities with bullet points
    st.markdown("""
        <div style='margin-left: 40px; margin-bottom: 25px; line-height: 1.6;'>
            <strong>Key Capabilities:</strong>
            <ul>
                <li><strong>Universal Document Support:</strong> Process and analyze PDF, TXT, and DOC files with automated content extraction</li>
                <li><strong>Dual Intelligence Systems:</strong> Choose between lightweight TF-IDF analysis or Mistral's powerful embedding model for optimal performance</li>
                <li><strong>Smart Content Processing:</strong> Leverage advanced chunking strategies - Fixed-Size with overlap for consistency, or Sentence-Based for semantic integrity</li>
                <li><strong>Enterprise-Ready Infrastructure:</strong> Built on AWS services for scalability, security, and reliability</li>
                <li><strong>Interactive Query System:</strong> Get context-aware responses powered by Mistral's language models (open-mistral-7b, mistral-small-latest, codestral-latest)</li>
                <li><strong>Real-time Diagnostics:</strong> Monitor system health and processing status with built-in diagnostic tools</li>
                <li><strong>Customizable Response Generation:</strong> Fine-tune output with adjustable temperature and token controls</li>
            </ul>
        </div>
    """, unsafe_allow_html=True)
    
    st.divider()
    
    # Add tabs for different sections
    tab1, tab2, tab3 = st.tabs(["Document Processing", "Query Interface", "Analytics"])
    
    with tab1:
        left_col, right_col = st.columns([1, 2])
        with left_col:
            st.markdown("### Processing Configuration")
            
            vectorizer_type = st.selectbox(
                "Choose Vectorizer or Embed Model:",
                ["TF-IDF", "Mistral-Embed"]
            )
            
            chunking_strategy = st.selectbox(
                "Choose Chunking Strategy:",
                ["Fixed-Size", "Sentence-Based"],
                help="Fixed-Size: Overlapping chunks of consistent size\nSentence-Based: Chunks that preserve sentence boundaries"
            )
            
            # Initialize vectorizer
            global vectorizer
            vectorizer, is_fitted = initialize_vectorizer(vectorizer_type)

            if vectorizer_type == "TF-IDF" and not is_fitted:
                st.warning("⚠️ Not initialized!")
            
            # Document processing controls
            process_button = st.button("Embed Documents", use_container_width=True)
            diagnose_button = st.button("Run Diagnostics", use_container_width=True)
            
            st.markdown("### Generation Settings")
            model = st.selectbox(
                "Choose a Model:",
                ["open-mistral-7b", "mistral-small-latest", "codestral-latest"]
            )
            
            col1, col2 = st.columns(2)
            with col1:
                temperature = st.slider(
                    "Temperature", 
                    min_value=0.0, 
                    max_value=1.5, 
                    value=0.7, 
                    step=0.1
                )
            with col2:
                max_tokens = st.slider(
                    "Max Tokens", 
                    min_value=1, 
                    max_value=1000, 
                    value=980, 
                    step=1
                )
            file_uploader_ui()
            
        with right_col:
            if process_button:
                progress_container = st.container()
                with progress_container:
                    with st.spinner("Processing documents..."):
                        try:
                            # List all objects in S3 bucket root
                            objects = s3_client.list_objects_v2(Bucket=S3_BUCKET_NAME)
                            
                            # Filter for supported document types
                            supported_files = [
                                obj for obj in objects.get('Contents', [])
                                if obj['Key'].lower().endswith(('.pdf', '.txt', '.doc', '.docx'))
                            ]
                            
                            total_docs = len(supported_files)
                            
                            if total_docs > 0:
                                st.info(f"Found {total_docs} documents to process")
                                
                                for obj in supported_files:
                                    document_key = obj['Key']
                                    st.text(f"Processing: {document_key}")
                                    
                                    # Load and process document
                                    document_content = load_document_from_s3(S3_BUCKET_NAME, document_key)
                                    
                                    if document_content:
                                        # Chunk document and store embeddings
                                        chunks = chunk_document(document_content, chunking_strategy)
                                        if store_embeddings(chunks, vectorizer_type):
                                            st.success(f"Successfully processed {document_key}")
                                        else:
                                            st.error(f"Failed to process {document_key}")
                                    else:
                                        st.error(f"Could not load content from {document_key}")
                                
                                st.success("Document processing completed!")
                            else:
                                st.warning("No supported documents found in the bucket")
                                
                        except Exception as e:
                            st.error(f"Error during document processing: {str(e)}")
                            logger.error(f"Document processing error: {str(e)}")

            if diagnose_button:
                diagnose_document_processing(vectorizer_type)
    
    with tab2:
        # Query interface
        st.markdown("<div style='margin-top: 30px;'>", unsafe_allow_html=True)
        prompt = st.text_area("Enter your prompt:", height=100)
        col1, col2 = st.columns([1, 5])
        with col1:
            submit_button = st.button("Submit")
            
        if submit_button:
            if vectorizer_type == "TF-IDF" and not hasattr(vectorizer, 'vocabulary_'):
                st.error("Please embed documents first!")
                return
                
            if prompt:
                # Check cache first
                cached_response = get_cached_response(prompt, "")
                if cached_response:
                    st.success("Retrieved from cache!")
                    st.write(cached_response)
                else:
                    with st.spinner("Processing query..."):
                        relevant_chunks = retrieve_relevant_chunks(prompt, vectorizer_type)
                        
                        if relevant_chunks:
                            # Show top 2 similarity scores
                            st.markdown("### Top Matching Scores:")
                            for i, chunk_info in enumerate(relevant_chunks[:2]):
                                st.markdown(f"**Match {i+1}:** {chunk_info['similarity']:.4f}")
                            
                            # Expandable context
                            with st.expander("View Complete Context", expanded=False):
                                for chunk_info in relevant_chunks:
                                    st.info(f"Similarity: {chunk_info['similarity']:.4f}")
                                    st.write(chunk_info['chunk'])
                                    st.divider()

                            # Generate and show response
                            context = " ".join(chunk["chunk"] for chunk in relevant_chunks)
                            combined_prompt = f"Context: {context}\n\nQuestion: {prompt}\n\nPlease provide a detailed answer based on the context above."
                            
                            with st.spinner("Generating response..."):
                                response = call_mistral_api(combined_prompt, model, temperature, max_tokens)
                                st.markdown("### Response:")
                                st.write(response)
                                
                                # Store in cache and history
                                store_cached_response(prompt, "", response)
                                st.session_state.history.append({
                                    'timestamp': datetime.now().isoformat(),
                                    'prompt': prompt,
                                    'response': response,
                                    'model': model
                                })
                        else:
                            st.warning("No relevant content found.")
            else:
                st.warning("Please enter a prompt.")
    
    with tab3:
        st.subheader("Analytics")
        
        # Show query history
        if st.session_state.history:
            st.write("Recent Queries:")
            for item in reversed(st.session_state.history[-5:]):
                with st.expander(f"{item['timestamp']} - {item['prompt'][:50]}..."):
                    st.write(f"**Prompt:** {item['prompt']}")
                    st.write(f"**Response:** {item['response']}")
                    st.write(f"**Model:** {item['model']}")
        
        # Visualize embeddings if available
        if 'embeddings' in st.session_state:
            st.plotly_chart(
                visualize_embeddings(
                    st.session_state.embeddings,
                    st.session_state.get('chunk_labels')
                )
            )

if __name__ == "__main__":
    main()